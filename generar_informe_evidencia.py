#!/usr/bin/env python3
"""Genera un informe DOCX presentable desde la evidencia JSON del auditor.

Uso:
    python generar_informe_evidencia.py evidencia.json informe.docx

El JSON conserva la evidencia canónica y el DOCX aporta una lectura ejecutiva,
trazabilidad, detalle de hallazgos y una matriz técnica para revisión humana.
"""

from __future__ import annotations

import argparse
import hashlib
import json
import re
from collections import Counter
from datetime import datetime
from pathlib import Path
from typing import Any, Iterable, Sequence

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor, Twips


PAGE_WIDTH_DXA = 12240
CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGINS_DXA = {"top": 80, "bottom": 80, "start": 120, "end": 120}

NAVY = "1F4E78"
BLUE = "2E74B5"
LIGHT_BLUE = "D9EAF7"
PALE_BLUE = "EDF4FA"
PALE_GREEN = "E2F0D9"
PALE_AMBER = "FFF2CC"
PALE_RED = "FCE4D6"
PALE_GRAY = "F2F4F7"
MID_GRAY = "667085"
DARK = "1D2939"
WHITE = "FFFFFF"
BORDER = "D0D5DD"

SEVERITY_COLORS = {
    "CRITICA": ("8B1E1E", "FFFFFF"),
    "ALTA": ("C65911", "FFFFFF"),
    "MEDIA": ("FFD966", "1D2939"),
    "BAJA": ("A9D18E", "1D2939"),
    "INFORMATIVA": ("D9EAF7", "1D2939"),
}

PILLAR_NAMES = {
    1: "Identidad y control de acceso",
    2: "Configuración, límites y consumo",
    3: "Integridad de dependencias y cadena de suministro",
}

PILLAR_SECTION_NAMES = {
    1: "Identidad y control de acceso",
    2: "Configuración, límites y consumo",
    3: "Integridad y cadena de suministro",
}

PILLAR_RISK_PROFILES = {
    1: (
        "quiebre de las fronteras de identidad, propiedad y rol",
        "Restablecer la autorización de mínimo privilegio y repetir las pruebas con las dos identidades ficticias.",
    ),
    2: (
        "configuraciones permisivas y controles de consumo incompletos",
        "Endurecer la configuración de despliegue y validar límites, cabeceras y orígenes con el mismo presupuesto.",
    ),
    3: (
        "falta de reproducibilidad y garantías de integridad en dependencias, CI/CD y registros",
        "Fijar referencias, incorporar lockfiles y hashes, y verificar nuevamente la cadena de integridad.",
    ),
}

FINDING_IMPACTS = {
    "P1-AGENT-004": "La condición observada puede permitir que una identidad de bajo privilegio opere mediante el agente con facultades superiores y acceda a herramientas o datos fuera de su rol.",
    "P1-SCOPE-006": "El agente actúa como canal de acceso indebido: una identidad de analista obtiene, a través del chat, más solicitudes de las que la API directa le entrega, lo que confirma con datos observables que las herramientas se ejecutan con privilegios ajenos al solicitante.",
    "P1-BOLA-002": "La autorización por objeto es insuficiente: un usuario autenticado podría consultar recursos pertenecientes a otra identidad, afectando la confidencialidad de los trámites.",
    "P1-ROLE-003": "El acceso de un analista a funciones administrativas puede exponer información sensible y habilitar acciones reservadas a coordinación.",
    "P1-BRUTE-005": "La ausencia de limitación observable aumenta la viabilidad de ataques automatizados contra credenciales y dificulta detectar intentos reiterados.",
    "P2-CORS-002": "Un origen no confiable podría realizar solicitudes autenticadas desde el navegador y leer respuestas protegidas si la sesión de la víctima está activa.",
    "P2-DEBUG-005": "Una configuración DEBUG habilitada puede revelar detalles internos, rutas, variables o trazas útiles para ampliar un ataque.",
    "P2-RATE-001": "Si se confirma, el indicador urgente permitiría eludir el presupuesto del agente y elevar el consumo de recursos; la evidencia actual todavía es inconclusa.",
    "P2-SECRET-004": "Una clave de respaldo pública o predecible reduce la confianza en firmas, sesiones o mecanismos criptográficos que dependan de ese secreto.",
    "P2-BODY-006": "La falta de un límite global de cuerpo puede facilitar solicitudes desproporcionadas y presión innecesaria sobre memoria, CPU o almacenamiento.",
    "P2-HEADERS-003": "La ausencia de cabeceras defensivas deja al navegador sin restricciones adicionales frente a interpretación de contenido, framing u otras superficies web.",
    "P3-AUDIT-001": "Sin encadenamiento verificable, una alteración o eliminación de eventos de auditoría puede no ser detectable durante una investigación posterior.",
    "R-A03-001": "Una versión móvil o no fijada impide reconstruir con certeza el artefacto evaluado y puede incorporar cambios no revisados en futuras instalaciones.",
    "R-A03-008": "Una referencia móvil de Git o de una acción CI/CD puede cambiar sin modificar el flujo versionado, debilitando la procedencia y la reproducibilidad del pipeline.",
    "R-A03-002": "La ausencia de lockfile permite resoluciones distintas entre ambientes y reduce la capacidad de reproducir, revisar y responder ante dependencias comprometidas.",
    "R-A08-001": "Sin un hash esperado no existe una comprobación determinista de que el artefacto consumido coincide con el artefacto aprobado.",
}

FINDING_CLOSURE_TESTS = {
    "P1-AGENT-004": "Repetir la regla con identidad de analista; la identidad efectiva debe conservar el rol autorizado, impedir herramientas privilegiadas y producir PASS.",
    "P1-SCOPE-006": "Repetir la misma tarea con la identidad de analista; la cantidad devuelta por la herramienta del agente debe coincidir con la de la API directa para esa cuenta y la regla debe producir PASS.",
    "P1-BOLA-002": "Repetir el acceso cruzado con las identidades A y B; el recurso ajeno debe responder 403/404 y la regla debe producir PASS.",
    "P1-ROLE-003": "Solicitar nuevamente el endpoint administrativo con el rol de analista; debe responder 403/404 y producir PASS.",
    "P1-BRUTE-005": "Repetir la secuencia acotada de intentos fallidos; debe observarse bloqueo, espera o limitación antes de agotar el presupuesto.",
    "P2-CORS-002": "Repetir con el origen no confiable; la respuesta no debe reflejarlo junto con credenciales y la regla debe producir PASS.",
    "P2-DEBUG-005": "Auditar la plantilla de despliegue corregida; DEBUG debe estar deshabilitado por defecto y la regla debe producir PASS.",
    "P2-RATE-001": "Completar la precondición y repetir con el mismo presupuesto; la regla debe finalizar sin ERROR y demostrar que el indicador no evade el límite.",
    "P2-SECRET-004": "Eliminar el valor de respaldo, inyectar un secreto administrado y repetir el análisis hasta obtener PASS sin exponer el secreto en la evidencia.",
    "P2-BODY-006": "Enviar el cuerpo de prueba acotado; la aplicación debe rechazarlo con el código configurado antes de procesarlo por completo.",
    "P2-HEADERS-003": "Repetir la solicitud HTTP y comprobar la presencia y el valor esperado de cada cabecera declarada en la política.",
    "P3-AUDIT-001": "Generar una cadena nueva, verificarla y alterar una copia controlada; la cadena original debe validar y la copia modificada debe ser detectada.",
    "R-A03-001": "Fijar la versión o referencia, regenerar el archivo reproducible y repetir la regla hasta obtener PASS.",
    "R-A03-008": "Sustituir la referencia móvil por un commit SHA inmutable y repetir el análisis hasta obtener PASS.",
    "R-A03-002": "Generar y versionar el lockfile correspondiente; repetir la regla y confirmar una resolución reproducible con PASS.",
    "R-A08-001": "Registrar el SHA-256 esperado, verificarlo durante el consumo y repetir la regla hasta obtener PASS.",
}

REPORT_AUTHORS = (
    "Diego Andrés García Álvarez",
    "Angel Danilo Marin Giraldo",
    "Brayan Alexander Salazar Reyes",
)
REPORT_AUTHOR_ROLE = "Ingenieros de Sistemas"


def _set_font(run, name: str = "Calibri", size: float | None = None, bold: bool | None = None,
              color: str | None = None, italic: bool | None = None) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def _set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def _set_cell_margins(cell, margins: dict[str, int] | None = None) -> None:
    values = dict(CELL_MARGINS_DXA)
    if margins:
        values.update(margins)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side in ("top", "bottom", "start", "end"):
        item = tc_mar.find(qn(f"w:{side}"))
        if item is None:
            item = OxmlElement(f"w:{side}")
            tc_mar.append(item)
        item.set(qn("w:w"), str(values[side]))
        item.set(qn("w:type"), "dxa")


def _set_table_borders(table, color: str = BORDER, size: str = "4") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), size)
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), color)


def apply_table_geometry(table, widths: Sequence[int]) -> None:
    widths = [int(width) for width in widths]
    if sum(widths) != CONTENT_WIDTH_DXA:
        raise ValueError(f"las columnas deben sumar {CONTENT_WIDTH_DXA} DXA: {widths}")
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(CONTENT_WIDTH_DXA))

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for col_index, width in enumerate(widths):
        table.columns[col_index].width = Twips(width)
    for row in table.rows:
        row.height = None
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        tr_pr = row._tr.get_or_add_trPr()
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)
        if len(row.cells) != len(widths):
            raise ValueError("la geometría exige filas sin celdas combinadas")
        for index, cell in enumerate(row.cells):
            cell.width = Twips(widths[index])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(widths[index]))
            _set_cell_margins(cell)
    _set_table_borders(table)


def _paragraph_border(paragraph, *, bottom_color: str = BLUE, bottom_size: str = "18") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), bottom_size)
    bottom.set(qn("w:space"), "8")
    bottom.set(qn("w:color"), bottom_color)
    p_bdr.append(bottom)


def _set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def _set_cell_text(cell, text: Any, *, bold: bool = False, color: str = DARK,
                   size: float = 9.0, name: str = "Calibri", align=None) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    if align is not None:
        paragraph.alignment = align
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    run = paragraph.add_run(str(text))
    _set_font(run, name=name, size=size, bold=bold, color=color)


def _add_table(document, rows: Iterable[Sequence[Any]], widths: Sequence[int],
               headers: Sequence[str] | None = None, header_fill: str = NAVY,
               font_size: float = 9.0):
    data = list(rows)
    table = document.add_table(rows=1 if headers else 0, cols=len(widths))
    if headers:
        for index, value in enumerate(headers):
            _set_cell_text(table.rows[0].cells[index], value, bold=True, color=WHITE, size=8.8)
            _set_cell_shading(table.rows[0].cells[index], header_fill)
        _set_repeat_table_header(table.rows[0])
    for row_values in data:
        row = table.add_row()
        for index, value in enumerate(row_values):
            _set_cell_text(row.cells[index], value, size=font_size)
            if len(table.rows) % 2 == 1:
                _set_cell_shading(row.cells[index], "F8FAFC")
    apply_table_geometry(table, widths)
    document.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def _add_label_value(document, label: str, value: Any) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.paragraph_format.keep_together = True
    run = paragraph.add_run(label + " ")
    _set_font(run, bold=True, color=NAVY, size=9.5)
    value_run = paragraph.add_run(str(value))
    _set_font(value_run, color=DARK, size=9.5)


def _add_callout(document, title: str, body: str, fill: str = LIGHT_BLUE,
                 accent: str = BLUE) -> None:
    table = document.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    _set_cell_shading(cell, fill)
    _set_cell_margins(cell, {"top": 170, "bottom": 170, "start": 200, "end": 200})
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(title)
    _set_font(run, bold=True, color=accent, size=11)
    paragraph = cell.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.08
    run = paragraph.add_run(body)
    _set_font(run, color=DARK, size=10)
    apply_table_geometry(table, [CONTENT_WIDTH_DXA])
    document.add_paragraph().paragraph_format.space_after = Pt(0)


def _add_page_number(paragraph) -> None:
    run = paragraph.add_run("Página ")
    _set_font(run, size=8.5, color=MID_GRAY)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, text, end])


def _configure_document(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    document.settings.odd_and_even_pages_header_footer = False
    section.different_first_page_header_footer = False

    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(DARK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, NAVY, 8, 4),
    ):
        style = document.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        style = document.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.167

    def populate_footer(footer: Any) -> None:
        paragraph = footer.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        _add_page_number(paragraph)

    populate_footer(section.footer)


def _clean_text(value: Any) -> str:
    if value is None:
        return "No aplica"
    if isinstance(value, bool):
        return "Sí" if value else "No"
    if isinstance(value, (dict, list)):
        return json.dumps(value, ensure_ascii=False, sort_keys=True)
    return str(value)


def _human_timestamp(value: str | None) -> str:
    if not value:
        return "No registrado"
    try:
        parsed = datetime.fromisoformat(value.replace("Z", "+00:00"))
        return parsed.strftime("%Y-%m-%d %H:%M:%S %z")
    except ValueError:
        return value


def _wrap_digest(value: str | None) -> str:
    if not value:
        return "No registrado"
    if ":" in value:
        prefix, digest = value.split(":", 1)
        return f"{prefix}:\n" + "\n".join(digest[i:i + 32] for i in range(0, len(digest), 32))
    return "\n".join(value[i:i + 32] for i in range(0, len(value), 32))


def _location(finding: dict[str, Any]) -> str:
    if finding.get("endpoint"):
        return finding["endpoint"]
    path = finding.get("archivo") or "repositorio"
    start = (finding.get("ubicacion") or {}).get("linea_inicio")
    return f"{path}:{start}" if start else path


def _evidence_lines(finding: dict[str, Any], limit: int = 4, max_chars: int = 420) -> list[str]:
    lines = []
    evidence = list(finding.get("evidencia", []))
    for item in evidence[:limit]:
        kind = str(item.get("tipo", "evidencia")).replace("_", " ")
        value = item.get("valor")
        if isinstance(value, list):
            rendered = ", ".join(_clean_text(element) for element in value) or "lista vacía"
        elif isinstance(value, dict):
            rendered = "; ".join(f"{key}={_clean_text(val)}" for key, val in value.items())
        else:
            rendered = _clean_text(value)
        source = item.get("fuente")
        if source:
            rendered += f" (fuente: {source})"
        if len(rendered) > max_chars:
            rendered = rendered[:max_chars - 3].rstrip() + "… [resumen; consulte el JSON canónico]"
        lines.append(f"{kind}: {rendered}")
    if len(evidence) > limit:
        lines.append(f"{len(evidence) - limit} elemento(s) adicional(es) conservado(s) en el JSON canónico.")
    return lines or ["La regla no registró valores adicionales."]


def _finding_impact(finding: dict[str, Any]) -> str:
    rule = str(finding.get("regla_id", ""))
    if rule in FINDING_IMPACTS:
        return FINDING_IMPACTS[rule]
    pillar = int(finding.get("pilar", 0) or 0)
    fallback = {
        1: "La condición puede debilitar una frontera de identidad o autorización y permitir operaciones fuera del rol esperado.",
        2: "La condición puede ampliar la superficie de ataque o permitir un consumo y una configuración distintos de los previstos.",
        3: "La condición reduce la reproducibilidad o la capacidad de demostrar la integridad y procedencia de los componentes evaluados.",
    }
    return fallback.get(pillar, "La condición observada requiere análisis de impacto dentro del contexto operativo del sistema.")


def _closure_test(finding: dict[str, Any]) -> str:
    rule = str(finding.get("regla_id", ""))
    if rule in FINDING_CLOSURE_TESTS:
        return FINDING_CLOSURE_TESTS[rule]
    if finding.get("estado_final") == "REQUIERE_REVISION":
        return "Completar la precondición indicada y repetir la misma regla hasta obtener una observación concluyente, sin ERROR."
    return f"Aplicar la corrección y repetir {rule or 'la regla'} con el mismo instrumento; el criterio determinista debe producir PASS."


def _treatment_priority(finding: dict[str, Any]) -> tuple[str, str]:
    state = str(finding.get("estado_final", "REQUIERE_REVISION"))
    severity = str(finding.get("severidad", "INFORMATIVA"))
    if state == "REQUIERE_REVISION":
        return "VALIDACIÓN", "Completar y repetir el control antes de presentar la condición como vulnerabilidad confirmada."
    if severity == "CRITICA":
        return "INMEDIATA", "Corregir antes de autorizar el siguiente despliegue y verificar el cierre con una corrida comparable."
    if severity == "ALTA":
        return "ALTA", "Asignar responsable y fecha antes del siguiente release; repetir la regla después de la corrección."
    if severity == "MEDIA":
        return "PLANIFICADA", "Incluir en el ciclo vigente de endurecimiento y verificar antes de cerrar el hallazgo."
    return "SEGUIMIENTO", "Registrar la decisión de tratamiento y comprobarla en la siguiente corrida programada."


def _add_findings_overview(document: Document, findings: Sequence[dict[str, Any]]) -> None:
    rows = []
    for pillar in (1, 2, 3):
        items = [item for item in findings if int(item.get("pilar", 0) or 0) == pillar]
        states = Counter(str(item.get("estado_final", "REQUIERE_REVISION")) for item in items)
        severities = Counter(str(item.get("severidad", "INFORMATIVA")) for item in items)
        profile, _ = PILLAR_RISK_PROFILES[pillar]
        rows.append((
            f"Pilar {pillar}",
            len(items),
            f"{states['CONFIRMADO']} conf. / {states['REQUIERE_REVISION']} rev.",
            f"{severities['CRITICA']} C · {severities['ALTA']} A · {severities['MEDIA']} M",
            profile if items else "Sin hallazgos promovidos en la corrida.",
        ))
    table = _add_table(
        document,
        rows,
        [1100, 700, 1600, 1700, 4260],
        headers=("Cobertura", "Total", "Estado", "Severidad", "Exposición principal"),
        font_size=8.5,
    )
    for row in table.rows[1:]:
        for index in (1, 2, 3):
            row.cells[index].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER


def _add_pillar_profile(document: Document, pillar: int, findings: Sequence[dict[str, Any]]) -> None:
    states = Counter(str(item.get("estado_final", "REQUIERE_REVISION")) for item in findings)
    severities = Counter(str(item.get("severidad", "INFORMATIVA")) for item in findings)
    profile, decision = PILLAR_RISK_PROFILES[pillar]
    _add_label_value(
        document,
        "Perfil de exposición:",
        f"{len(findings)} hallazgo(s): {states['CONFIRMADO']} confirmado(s), "
        f"{states['REQUIERE_REVISION']} en revisión; {severities['CRITICA']} crítico(s), "
        f"{severities['ALTA']} alto(s) y {severities['MEDIA']} medio(s). Foco: {profile}.",
    )
    _add_label_value(document, "Decisión de tratamiento:", decision)


def _add_finding(document, finding: dict[str, Any], number: int) -> None:
    severity = finding.get("severidad", "INFORMATIVA")
    state = finding.get("estado_final", "REQUIERE_REVISION")
    rule = finding.get("regla_id", "SIN-REGLA")
    title = finding.get("hallazgo", "Hallazgo sin título")
    pillar = int(finding.get("pilar", 0) or 0)
    confidence = finding.get("confianza", "NO REGISTRADA")
    priority, priority_detail = _treatment_priority(finding)

    heading = document.add_paragraph(style="Heading 2")
    heading.paragraph_format.keep_with_next = True
    run = heading.add_run(f"H-{number:02d} · {title}")
    _set_font(run, bold=True, color=NAVY, size=12.5)

    metadata = document.add_table(rows=1, cols=6)
    labels = (
        ("Regla", rule),
        ("Pilar", f"Pilar {pillar}"),
        ("Severidad", severity),
        ("Estado", state),
        ("Confianza", confidence),
        ("Prioridad", priority),
    )
    data_row = metadata.rows[0]
    for index, (label, value) in enumerate(labels):
        cell = data_row.cells[index]
        cell.text = ""
        label_paragraph = cell.paragraphs[0]
        label_paragraph.paragraph_format.space_after = Pt(1)
        label_run = label_paragraph.add_run(label.upper())
        _set_font(label_run, bold=True, color=MID_GRAY, size=7.0)
        value_paragraph = cell.add_paragraph()
        value_paragraph.paragraph_format.space_after = Pt(0)
        if index in (1, 2, 3, 4):
            label_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            value_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        value_run = value_paragraph.add_run(str(value))
        _set_font(value_run, bold=index in (0, 2, 3, 5), color=DARK, size=8.1)
    apply_table_geometry(metadata, [1450, 900, 1250, 1750, 1300, 2710])
    severity_fill, severity_text = SEVERITY_COLORS.get(str(severity), (PALE_GRAY, DARK))
    _set_cell_shading(data_row.cells[2], severity_fill)
    for run in data_row.cells[2].paragraphs[0].runs:
        _set_font(run, bold=True, color=severity_text, size=7.0)
    for run in data_row.cells[2].paragraphs[1].runs:
        _set_font(run, bold=True, color=severity_text, size=8.2)
    if state == "REQUIERE_REVISION":
        _set_cell_shading(data_row.cells[3], PALE_AMBER)
    if priority == "INMEDIATA":
        _set_cell_shading(data_row.cells[5], PALE_RED)

    _add_label_value(document, "Ubicación técnica:", _location(finding))
    _add_label_value(document, "Descripción técnica:", finding.get("detalle", "Sin detalle"))
    _add_label_value(document, "Impacto potencial:", _finding_impact(finding))
    _add_label_value(document, "Anclaje de referencia:", finding.get("categoria_owasp", "No registrado"))
    _add_label_value(document, "Evidencia decisiva:", "")
    for line in _evidence_lines(finding):
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.keep_together = True
        run = paragraph.add_run(line)
        _set_font(run, name="Consolas" if "sha256:" in line else "Calibri", size=9.1, color=DARK)
    _add_label_value(document, "Tratamiento recomendado:", finding.get("recomendacion", "No registrado"))
    _add_label_value(document, "Decisión de tratamiento:", priority_detail)
    _add_label_value(document, "Prueba de cierre:", _closure_test(finding))
    note = document.add_paragraph()
    note.paragraph_format.space_after = Pt(8)
    note.paragraph_format.keep_together = True
    r = note.add_run(
        f"Conclusión técnica: {state} con confianza {confidence}. La evidencia satisface el criterio "
        "determinista de la regla dentro del alcance ejecutado; no demuestra explotabilidad fuera de ese alcance."
        if state == "CONFIRMADO"
        else f"Conclusión técnica: {state} con confianza {confidence}. No debe presentarse como vulnerabilidad confirmada hasta repetir el control y resolver la causa de la evidencia inconclusa."
    )
    _set_font(r, italic=True, color="7F6000" if state == "REQUIERE_REVISION" else MID_GRAY, size=8.8)
    _paragraph_border(note, bottom_color=BORDER, bottom_size="4")


def _aggregate_rules(report: dict[str, Any]) -> list[tuple[str, int, int, int, int]]:
    rows = []
    for rule, counts in sorted(report.get("resumen", {}).get("eventos_por_regla", {}).items()):
        rows.append((
            rule,
            int(counts.get("PASS", 0)),
            int(counts.get("FAIL", 0)),
            int(counts.get("ERROR", 0)),
            int(counts.get("SKIP_JUSTIFICADO", 0)),
        ))
    return rows


def _risk_assessment(summary: dict[str, Any], findings: Sequence[dict[str, Any]]) -> tuple[str, str, str]:
    severities = summary.get("por_severidad", {})
    states = summary.get("por_estado", {})
    critical = int(severities.get("CRITICA", 0))
    high = int(severities.get("ALTA", 0))
    medium = int(severities.get("MEDIA", 0))
    confirmed = int(states.get("CONFIRMADO", 0))
    review = int(states.get("REQUIERE_REVISION", 0))
    if critical:
        level = "CRÍTICO"
        accent = "8B1E1E"
        fill = PALE_RED
        decision = (
            "La exposición observada requiere tratamiento prioritario antes de autorizar un despliegue. "
            "Los hallazgos críticos deben remediarse y verificarse nuevamente con el mismo instrumento."
        )
    elif high:
        level = "ALTO"
        accent = "C65911"
        fill = PALE_AMBER
        decision = (
            "La exposición observada requiere un plan de remediación con responsables y fechas, seguido "
            "de una nueva corrida comparable antes de aceptar el riesgo residual."
        )
    elif medium:
        level = "MEDIO"
        accent = "7F6000"
        fill = PALE_AMBER
        decision = "La exposición debe corregirse dentro del ciclo de ingeniería y confirmarse mediante una nueva corrida."
    else:
        level = "BAJO SEGÚN EL ALCANCE EJECUTADO"
        accent = "2F6B3B"
        fill = PALE_GREEN
        decision = "No se promovieron riesgos medios, altos o críticos dentro de las reglas y el alcance ejecutados."
    body = (
        f"Nivel técnico: {level}. La corrida registró {len(findings)} hallazgos, de los cuales "
        f"{confirmed} quedaron confirmados y {review} requieren revisión. Distribución relevante: "
        f"{critical} críticos, {high} altos y {medium} medios. {decision} Este dictamen se limita al "
        "snapshot, la instancia, las identidades y las reglas documentadas en la evidencia."
    )
    return body, fill, accent


def _priority_findings(findings: Sequence[dict[str, Any]], limit: int = 5) -> list[dict[str, Any]]:
    order = {"CRITICA": 0, "ALTA": 1, "MEDIA": 2, "BAJA": 3, "INFORMATIVA": 4}
    confirmed = [item for item in findings if item.get("estado_final") == "CONFIRMADO"]
    return sorted(
        confirmed,
        key=lambda item: (order.get(str(item.get("severidad")), 99), str(item.get("regla_id", ""))),
    )[:limit]


def _add_comparative_design(document: Document, report: dict[str, Any]) -> None:
    document.add_heading("3. Diseño comparativo A/B", level=1)
    comparison = report.get("comparacion_ab") or {}
    condition = (report.get("corrida") or {}).get("condicion", "No registrada")
    instrument = report.get("diseno_comparativo") or {}
    if not comparison:
        document.add_paragraph(
            f"La evidencia corresponde a la condición {condition}. Las condiciones A y B ejecutan las "
            "mismas reglas; para calcular el efecto de la intervención debe generar la otra condición y "
            "usar --comparar-con."
        )
        _add_label_value(document, "Hash del instrumento:", instrument.get("instrumento_sha256", "No registrado"))
        return

    changed_rules = list(comparison.get("reglas_con_cambio") or [])
    comparison_rows = [
        ("Corrida A", comparison.get("condicion_inicial_id", "No registrada")),
        ("Corrida B", comparison.get("intervencion_id", "No registrada")),
        ("Instrumento idéntico", _clean_text(comparison.get("instrumento_identico"))),
        ("Política idéntica", _clean_text(comparison.get("politica_identica"))),
        ("Repositorio cambió", _clean_text(comparison.get("repositorio_cambio"))),
        ("Hallazgos A / B", f"{(comparison.get('hallazgos') or {}).get('A', 0)} / {(comparison.get('hallazgos') or {}).get('B', 0)}"),
        ("Reglas con cambio", ", ".join(changed_rules) if changed_rules else "Ninguna"),
    ]
    _add_table(document, comparison_rows, [2500, 6860], headers=("Criterio", "Resultado"), font_size=9.0)
    comparable = bool(comparison.get("instrumento_identico") and comparison.get("politica_identica"))
    if not comparable:
        _add_callout(
            document,
            "Comparación no válida",
            "El instrumento o la política difieren entre condiciones. No atribuya los cambios a la intervención hasta repetir ambas corridas con parámetros equivalentes.",
            fill=PALE_RED,
            accent="8B1E1E",
        )
    elif changed_rules:
        _add_callout(
            document,
            "Efecto observable de la intervención",
            f"La comparación es técnicamente válida y registró cambios en {len(changed_rules)} regla(s): {', '.join(changed_rules)}. Cada transición debe revisarse contra la evidencia de la regla.",
            fill=PALE_GREEN,
            accent="2F6B3B",
        )
    else:
        _add_callout(
            document,
            "Intervención sin cambio observable",
            "Las corridas son comparables, pero ninguna regla cambió de resultado. Con este instrumento no se observó reducción del riesgo; revise que la intervención haya quedado activa y repita la verificación.",
            fill=PALE_AMBER,
            accent="7F6000",
        )


def _add_rule_matrix(document, report: dict[str, Any]) -> None:
    document.add_heading("Anexo A. Matriz de resultados por regla", level=1)
    document.add_paragraph(
        "Cada evento conserva un resultado mecánico. FAIL promueve un hallazgo; PASS acredita que el control no detectó la condición; ERROR requiere revisión; SKIP_JUSTIFICADO documenta por qué una regla no se ejecutó."
    )
    rows = _aggregate_rules(report)
    table = _add_table(
        document,
        rows,
        [3000, 1300, 1300, 1300, 2460],
        headers=("Regla", "PASS", "FAIL", "ERROR", "SKIP justificado"),
        font_size=8.6,
    )
    for row in table.rows[1:]:
        for index in range(1, 5):
            row.cells[index].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER


def build_report(report: dict[str, Any], source_path: Path, output_path: Path) -> None:
    document = Document()
    _configure_document(document)
    document.core_properties.title = "Informe técnico de ciberseguridad — Pilares 1, 2 y 3"
    document.core_properties.subject = "Evaluación profesional y evidencia reproducible de controles de seguridad"
    document.core_properties.author = "; ".join(REPORT_AUTHORS)
    document.core_properties.keywords = "ciberseguridad, auditoría, OWASP, evidencia, Tramitia, pilares 1 2 3"

    corrida = report.get("corrida", {})
    summary = report.get("resumen", {})
    runtime = report.get("runtime", {})
    scope = report.get("alcance", {})
    findings = list(report.get("hallazgos", []))
    json_hash = "sha256:" + hashlib.sha256(source_path.read_bytes()).hexdigest()

    # Preset standard_business_brief con apertura memo_masthead.
    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run("EVALUACIÓN DE SEGURIDAD DE SOFTWARE")
    _set_font(r, name="Calibri", size=9, bold=True, color=BLUE)
    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run("INFORME TÉCNICO DE CIBERSEGURIDAD")
    _set_font(r, name="Calibri", size=23, bold=True, color=NAVY)
    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run("Evaluación comparativa de los pilares 1, 2 y 3")
    _set_font(r, size=13, color=MID_GRAY)

    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("ELABORADO POR")
    _set_font(r, size=8, bold=True, color=MID_GRAY)
    for author in REPORT_AUTHORS:
        p = document.add_paragraph()
        p.paragraph_format.space_after = Pt(1)
        r = p.add_run(author)
        _set_font(r, size=10.5, bold=True, color=DARK)
    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run(REPORT_AUTHOR_ROLE)
    _set_font(r, size=9.5, italic=True, color=BLUE)

    metadata_rows = [
        ("Sistema evaluado", findings[0].get("caso_id", "Aplicación autorizada") if findings else "Aplicación autorizada"),
        ("Corrida", corrida.get("corrida_id", "No registrada")),
        ("Fecha de ejecución", _human_timestamp(corrida.get("timestamp_utc"))),
        ("Alcance", f"{runtime.get('base_url', 'sin pruebas HTTP')} + repositorio local"),
        ("Estado global", f"{summary.get('por_estado', {}).get('CONFIRMADO', 0)} confirmados · {summary.get('por_estado', {}).get('REQUIERE_REVISION', 0)} requiere revisión"),
        ("Fuente canónica", source_path.name),
    ]
    table = document.add_table(rows=0, cols=2)
    for label, value in metadata_rows:
        row = table.add_row()
        _set_cell_text(row.cells[0], label.upper(), bold=True, color=MID_GRAY, size=7.8)
        _set_cell_text(row.cells[1], value, color=DARK, size=9.2)
    apply_table_geometry(table, [2100, 7260])
    _set_table_borders(table, color=WHITE, size="0")

    assessment, assessment_fill, assessment_accent = _risk_assessment(summary, findings)
    _add_callout(
        document,
        "Dictamen ejecutivo",
        assessment,
        fill=assessment_fill,
        accent=assessment_accent,
    )

    document.add_heading("1. Dictamen y prioridades de remediación", level=1)
    document.add_paragraph(
        "La priorización combina severidad, estado final y evidencia observable. No sustituye la decisión formal de aceptación de riesgo del responsable del sistema."
    )
    priorities = _priority_findings(findings)
    if priorities:
        for item in priorities:
            p = document.add_paragraph(style="List Number")
            r = p.add_run(
                f"[{item.get('severidad', 'INFORMATIVA')}] {item.get('regla_id', 'SIN-REGLA')} — "
                f"{item.get('hallazgo', 'Hallazgo sin título')}. Acción: {item.get('recomendacion', 'Revisar el control.')}"
            )
            _set_font(r, size=10.2, color=DARK)
    else:
        document.add_paragraph("No existen hallazgos confirmados para priorizar en esta corrida.")

    document.add_heading("2. Resumen de resultados", level=1)
    pillar_rows = []
    per_pillar = report.get("resumen_por_pilar", {})
    for pillar in (1, 2, 3):
        item = per_pillar.get(str(pillar), {})
        state = item.get("por_estado", {})
        severities = item.get("por_severidad", {})
        pillar_rows.append((
            f"Pilar {pillar}",
            PILLAR_NAMES[pillar],
            item.get("hallazgos_total", 0),
            state.get("CONFIRMADO", 0),
            state.get("REQUIERE_REVISION", 0),
            f"C {severities.get('CRITICA', 0)} · A {severities.get('ALTA', 0)} · M {severities.get('MEDIA', 0)}",
        ))
    _add_table(document, pillar_rows, [1200, 3200, 1100, 1200, 1360, 1300],
               headers=("Pilar", "Alcance", "Total", "Confirmado", "Revisión", "Severidad C/A/M"), font_size=8.5)

    document.add_paragraph(
        "Lectura recomendada: tratar primero los hallazgos confirmados críticos y altos; después, los medios. Todo resultado REQUIERE_REVISION debe repetirse con la precondición o el presupuesto indicado antes de presentarlo como vulnerabilidad."
    )

    _add_comparative_design(document, report)

    document.add_heading("4. Alcance y método", level=1)
    method_rows = [
        ("Pilar 1", "Pruebas HTTP autorizadas", "Autenticación, propiedad de objetos, límites de rol, identidad efectiva y resistencia a intentos fallidos."),
        ("Pilar 2", "Pruebas HTTP + patrones fuente", "Límites de consumo, CORS, cabeceras, secretos por defecto, modo DEBUG y tamaño de cuerpo."),
        ("Pilar 3", "Análisis estático + integridad", "Versiones, lockfiles, hashes, procedencia, referencias inmutables y cadena verificable del registro de auditoría."),
    ]
    _add_table(document, method_rows, [1500, 2600, 5260], headers=("Pilar", "Técnica", "Cobertura"), font_size=8.9)
    for text in (
        f"Se inventariaron {scope.get('archivos_inventariados', 0)} archivos ({scope.get('bytes_inventariados', 0):,} bytes) y {len(report.get('componentes', []))} componentes.",
        f"Se realizaron {runtime.get('peticiones_realizadas', 0)} peticiones HTTP contra {runtime.get('base_url', 'una instancia no registrada')}.",
        "Los criterios de decisión son deterministas: estados HTTP, valores JSON, cabeceras o patrones concretos; el informe no usa IA para clasificar cada control.",
        "Las credenciales se suministraron por variables de entorno y la evidencia fue redactada antes de escribirse.",
    ):
        p = document.add_paragraph(style="List Bullet")
        p.add_run(text)

    document.add_heading("5. Trazabilidad e integridad de la evidencia", level=1)
    integrity_rows = [
        ("Identificador de corrida", corrida.get("corrida_id", "No registrado")),
        ("Inicio / fin", f"{_human_timestamp(corrida.get('timestamp_utc'))} / {_human_timestamp(corrida.get('finalizada_utc'))}"),
        ("Duración", f"{corrida.get('duracion_segundos', 0)} segundos"),
        ("Hash del snapshot", _wrap_digest(corrida.get("repo_hash"))),
        ("Hash del JSON fuente", _wrap_digest(json_hash)),
        ("Hash de configuración", _wrap_digest((corrida.get("config_hashes") or {}).get("effective_config"))),
        ("Herramienta", f"{report.get('tool', {}).get('name', 'auditor')} {report.get('tool', {}).get('version', 'sin versión')}"),
        ("Modificación del repositorio", "No" if not report.get("tool", {}).get("repository_modified") else "Sí"),
    ]
    integrity_table = _add_table(document, integrity_rows, [2800, 6560], headers=("Dato de custodia", "Valor"), font_size=8.6)
    for row in integrity_table.rows[1:]:
        if "Hash" in row.cells[0].text:
            for run in row.cells[1].paragraphs[0].runs:
                _set_font(run, name="Consolas", size=8.0, color=DARK)

    document.add_heading("6. Hallazgos sustentados de los tres pilares", level=1)
    document.add_paragraph(
        "Esta sección desarrolla de forma explícita los pilares 1, 2 y 3. Cada hallazgo conecta la observación objetiva con su impacto potencial, prioridad de tratamiento, anclaje OWASP, evidencia decisiva y prueba de cierre. CONFIRMADO significa que la observación cumplió el criterio determinista de promoción; REQUIERE_REVISION indica que todavía no existe evidencia suficiente para una conclusión definitiva."
    )
    _add_findings_overview(document, findings)
    counter = 1
    for pillar in (1, 2, 3):
        pillar_findings = [item for item in findings if int(item.get("pilar", 0)) == pillar]
        count_label = "1 hallazgo" if len(pillar_findings) == 1 else f"{len(pillar_findings)} hallazgos"
        document.add_heading(
            f"Pilar {pillar}. {PILLAR_SECTION_NAMES[pillar]} ({count_label})",
            level=1,
        )
        _add_pillar_profile(document, pillar, pillar_findings)
        if not pillar_findings:
            document.add_paragraph("No se promovieron hallazgos para este pilar dentro de la corrida y el alcance documentados.")
        for finding in pillar_findings:
            _add_finding(document, finding, counter)
            counter += 1

    document.add_heading("7. Controles sin hallazgo y limitaciones", level=1)
    pass_results = [item for item in report.get("resultados_reglas", []) if item.get("resultado") == "PASS"]
    skip_results = [item for item in report.get("resultados_reglas", []) if item.get("resultado") == "SKIP_JUSTIFICADO"]
    pass_counts = Counter(item.get("regla_id", "SIN-REGLA") for item in pass_results)
    skip_counts = Counter(item.get("regla_id", "SIN-REGLA") for item in skip_results)
    if pass_counts:
        document.add_heading("Controles aprobados", level=2)
        for rule, count in sorted(pass_counts.items()):
            p = document.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.line_spacing = 1.05
            p.add_run(f"{rule}: {count} evento(s) PASS.")
    if skip_counts:
        document.add_heading("Reglas omitidas con justificación", level=2)
        for rule, count in sorted(skip_counts.items()):
            details = sorted({item.get("detalle", "sin detalle") for item in skip_results if item.get("regla_id") == rule})
            p = document.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.line_spacing = 1.05
            p.add_run(f"{rule}: {count} evento(s). {'; '.join(details)}")
    document.add_heading("Limitaciones declaradas", level=2)
    limitations = report.get("limitaciones")
    limitation_items = [limitations] if isinstance(limitations, str) else list(limitations or [])
    limitation_items.extend([
        "La prueba se limita a una instancia local y a un snapshot del repositorio; otros despliegues pueden producir resultados distintos.",
        "La ausencia de un hallazgo no demuestra ausencia de vulnerabilidades fuera de las reglas activas.",
        "Los resultados ERROR y REQUIERE_REVISION no deben presentarse como vulnerabilidades confirmadas sin repetir y completar el control.",
        "Los controles activos pueden modificar estado de prueba; deben ejecutarse sobre datos y ambientes autorizados.",
    ])
    for text in limitation_items:
        p = document.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.05
        p.add_run(text)

    _add_rule_matrix(document, report)
    document.add_page_break()
    heading = document.add_paragraph(style="Heading 1")
    heading.paragraph_format.keep_with_next = True
    run = heading.add_run("Procedimiento reproducible — Anexo B")
    _set_font(run, bold=True, color=BLUE, size=16)
    document.add_paragraph(
        "Ejemplo para Tramitia. Las credenciales reales se suministran en variables de entorno y no deben escribirse en el comando ni en el perfil versionado."
    )
    command = (
        "python auditor_tramitia.py --base-url http://127.0.0.1:5050 "
        "--repo <copia-local-autorizada> --config auditor_config.tramitia.example.json "
        "--pilares 1,2,3 --condicion B --comparar-con resultados/evidencia_A.json "
        "--autorizado --permitir-pruebas-activas --out resultados/evidencia_B.json "
        "--out-docx resultados/Informe_Profesional_Ciberseguridad_B.docx"
    )
    table = document.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    _set_cell_shading(cell, PALE_GRAY)
    _set_cell_text(cell, command, name="Consolas", size=8.4, color=DARK)
    apply_table_geometry(table, [CONTENT_WIDTH_DXA])
    document.add_paragraph(
        "Verificación de integridad: calcular SHA-256 del JSON y comparar el resultado con el hash consignado en la sección 5. El hash del snapshot identifica exactamente el conjunto de archivos examinado durante la corrida."
    )

    document.add_heading("Anexo C. Criterios de interpretación", level=1)
    criteria_rows = [
        ("CONFIRMADO", "La regla determinista observó la condición definida como fallo.", "Puede presentarse como evidencia del snapshot evaluado."),
        ("REQUIERE_REVISION", "La prueba fue inconclusa, insuficiente o requiere calibración.", "No presentarlo como falla confirmada; repetir el control."),
        ("PASS", "No se observó la condición bajo el criterio y alcance ejecutados.", "Acredita el control puntual, no la ausencia total de riesgo."),
        ("SKIP_JUSTIFICADO", "La regla no aplicó o faltó una fuente local necesaria.", "Conservar la justificación y, si procede, ampliar el alcance."),
    ]
    _add_table(document, criteria_rows, [2000, 3800, 3560], headers=("Estado", "Significado", "Uso correcto"), font_size=8.8)

    document.add_heading("Responsables del informe", level=1)
    document.add_paragraph(
        "El informe fue preparado mediante un instrumento determinista y revisado técnicamente por el siguiente equipo:"
    )
    author_rows = [(author, "Ingeniero de Sistemas") for author in REPORT_AUTHORS]
    _add_table(document, author_rows, [5600, 3760], headers=("Nombre", "Formación profesional"), font_size=9.2)

    output_path = output_path.expanduser().resolve()
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(description="Genera un informe DOCX desde la evidencia JSON del auditor")
    parser.add_argument("evidence", type=Path, help="archivo JSON generado por auditor_tramitia.py")
    parser.add_argument("output", type=Path, help="ruta del informe DOCX")
    args = parser.parse_args(argv)
    source = args.evidence.expanduser().resolve()
    if not source.is_file():
        parser.error(f"no existe la evidencia: {source}")
    try:
        report = json.loads(source.read_text(encoding="utf-8"))
    except (OSError, json.JSONDecodeError) as exc:
        parser.error(f"no se pudo leer la evidencia: {exc}")
    required = {"corrida", "alcance", "hallazgos", "resumen", "resultados_reglas"}
    missing = sorted(required.difference(report))
    if missing:
        parser.error("el JSON no tiene el esquema esperado; faltan: " + ", ".join(missing))
    build_report(report, source, args.output)
    print(args.output.expanduser().resolve())
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
